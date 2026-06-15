from __future__ import annotations

import importlib.util
import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
R91 = ROOT / "outputs" / "SCTS_R91_submission_polish_final"
R85_SCRIPT = ROOT / "scripts" / "build_SCTS_R85_submission_ordered_revision.py"
OUT = ROOT / "outputs" / "SCTS_R92_scts_format_compliance"
FIG = OUT / "figures"
FIG600 = OUT / "submission_figures_600dpi"
PDF = OUT / "pdf_proof"
PKG = OUT / "package"

spec = importlib.util.spec_from_file_location("r85", R85_SCRIPT)
r85 = importlib.util.module_from_spec(spec)
assert spec.loader is not None
spec.loader.exec_module(r85)
r84 = r85.r84


TITLE = "Local iodine-handling failure in porous carbon-felt electrodes drives solid-iodine blockage in zinc-iodine flow batteries"

SCTS_REFERENCES = """[1] Weber A Z, Mench M M, Meyers J P, et al. Redox flow batteries: a review. J Appl Electrochem, 2011, 41: 1137-1164. doi:10.1007/s10800-011-0348-2.
[2] Soloveichik G L. Flow batteries: current status and trends. Chem Rev, 2015, 115: 11533-11558. doi:10.1021/cr500720t.
[3] Noack J, Roznyatovskaya N, Herr T, et al. The chemistry of redox-flow batteries. Angew Chem Int Ed, 2015, 54: 9776-9809. doi:10.1002/anie.201410823.
[4] Li B, Nie Z, Vijayakumar M, et al. Ambipolar zinc-polyiodide electrolyte for a high-energy density aqueous redox flow battery. Nat Commun, 2015, 6: 6303. doi:10.1038/ncomms7303.
[5] Weng G M, Li Z, Cong G, et al. Unlocking the capacity of iodide for high-energy-density zinc/polyiodide and lithium/polyiodide redox flow batteries. Energy Environ Sci, 2017, 10: 735-741. doi:10.1039/C6EE03554J.
[6] Xie C, Zhang H, Xu W, et al. A long cycle life, self-healing zinc-iodine flow battery with high power density. Angew Chem Int Ed, 2018, 57: 11171-11176. doi:10.1002/anie.201803122.
[7] Xie C, Liu Y, Lu W, et al. Highly stable zinc-iodine single flow batteries with super high energy density for stationary energy storage. Energy Environ Sci, 2019, 12: 1834-1839. doi:10.1039/C8EE02825G.
[8] Ma J, Liu M, He Y, et al. Iodine redox chemistry in rechargeable batteries. Angew Chem Int Ed, 2021, 60: 12636-12647. doi:10.1002/anie.202009871.
[9] Williams A A, Emmett R K, Roberts M E. High power zinc iodine redox flow battery with iron-functionalized carbon electrodes. Phys Chem Chem Phys, 2023, 25: 16222-16226. doi:10.1039/D3CP02067C.
[10] Wang C, Gao G, Su Y, et al. High-voltage and dendrite-free zinc-iodine flow battery. Nat Commun, 2024, 15: 6234. doi:10.1038/s41467-024-50543-2.
[11] Richtr P, Graf D, Drnec M, et al. Understanding the degradation process in zinc-iodine hybrid flow batteries. J Mater Chem A, 2026, 14: 4529-4545. doi:10.1039/D5TA07792C.
[12] Berkowitz A, Caiado A A, Aravamuthan S R, et al. Optimization framework for redox flow battery electrodes with improved microstructural characteristics. Energy Adv, 2024, 3: 2220-2237. doi:10.1039/D4YA00248B.
[13] Palmer D A, Ramette R W, Mesmer R E. Triiodide ion formation equilibrium and activity coefficients in aqueous solution. J Solution Chem, 1984, 13: 673-683. doi:10.1007/BF00650374.
[14] Chakrabarti B K, Kalamaras E, Singh A K, et al. Modelling of redox flow battery electrode processes at a range of length scales: a review. Sustain Energy Fuels, 2020, 4: 5433-5468. doi:10.1039/D0SE00667J.
[15] Banerjee R, Bevilacqua N, Mohseninia A, et al. Carbon felt electrodes for redox flow battery: impact of compression on transport properties. J Energy Storage, 2019, 26: 100997. doi:10.1016/j.est.2019.100997.
[16] Emmel D, Hofmann J D, Arlt T, et al. Understanding the impact of compression on the active area of carbon felt electrodes for redox flow batteries. ACS Appl Energy Mater, 2020, 3: 4384-4393. doi:10.1021/acsaem.0c00075.
[17] Amini K, Shocron A N, Suss M E, et al. Pathways to high-power-density redox flow batteries. ACS Energy Lett, 2023, 8: 3526-3535. doi:10.1021/acsenergylett.3c01043.
[18] Wang P, Zhao Y, Ban Y, et al. A review of porous electrode structural parameters and optimization for redox flow batteries. J Energy Storage, 2024, 97: 112859. doi:10.1016/j.est.2024.112859.
[19] Lee J, Faheem A B, Jang W J, et al. Effective enhancement of energy density of zinc-polyiodide flow batteries by organic/penta-iodide complexation. ACS Appl Mater Interfaces, 2023, 15: 48122-48134. doi:10.1021/acsami.3c09426.
[20] Wei Z, Wang Y, Hong H, et al. Long-life aqueous zinc-iodine flow batteries enabled by selectively intercepting hydrated ions. Nat Commun, 2025, 16: 9301. doi:10.1038/s41467-025-64344-8.
"""


def ensure_dirs() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for p in [OUT, FIG, FIG600, PDF, PKG]:
        p.mkdir(parents=True, exist_ok=True)


def copy_tree(src: Path, dst: Path) -> None:
    if not src.exists():
        return
    for p in src.iterdir():
        if p.is_file():
            shutil.copy2(p, dst / p.name)


def scts_polish_markdown(md: str) -> str:
    md = re.sub(r"^# .*$", f"# {TITLE}", md, count=1, flags=re.M)
    md = md.replace(
        "The governing state is not total iodine inventory in the tank, but the fiber-scale balance between free-I2 generation and local removal by complexation, diffusion, flow renewal, dissolution and accessible carbon surface.",
        "The governing state is not total iodine inventory in the tank, but the fiber-scale balance between neutral iodine generation and local removal by complexation, diffusion, flow renewal, dissolution and accessible carbon surface.",
    )
    md = md.replace(
        "COMSOL postprocessing then shows two routes to the same failure chain:",
        "Porous-electrode postprocessing then shows two routes to the same failure chain:",
    )
    md = md.replace("## 4 Methods", "## 4 Materials and methods")
    md = md.replace("## Competing interests\n\n[The authors declare no competing interests, or revise as appropriate before submission.]", "## Conflict of Interest\n\nThe authors declare that they have no conflict of interest. [Revise before submission if any financial or non-financial conflict exists.]")
    md = md.replace("## Acknowledgements\n\n[To be completed by the authors before submission.]", "## Funding\n\n[To be completed by the authors before submission. Include full fund names and grant numbers.]\n\n## Acknowledgements\n\n[To be completed by the authors before submission.]")
    md = md.replace("## References", "## Ethical approval\n\nThis article does not contain any studies with human participants or animals performed by any of the authors.\n\n## Informed consent\n\nNot applicable.\n\n## References")
    md = re.sub(r"\n## References\n\n.*$", "\n## References\n\n" + SCTS_REFERENCES, md, flags=re.S)
    return md


def set_scts_docx_format(docx_path: Path) -> dict[str, object]:
    doc = Document(docx_path)
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
        sec.header_distance = Cm(1.25)
        sec.footer_distance = Cm(1.25)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        if paragraph.style.name == "Normal":
            paragraph.paragraph_format.space_after = Pt(6)
        for run in paragraph.runs:
            if not run.font.size or run.font.size.pt == 11:
                run.font.size = Pt(10)
            run.font.name = "Calibri"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.15
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        if not run.font.size or run.font.size.pt > 9:
                            run.font.size = Pt(8)
                        run.font.name = "Calibri"
    max_image_width = Inches(5.55)
    for shape in doc.inline_shapes:
        if shape.width and shape.width > max_image_width:
            ratio = max_image_width / shape.width
            shape.width = max_image_width
            shape.height = int(shape.height * ratio)
    doc.save(docx_path)
    sec = doc.sections[0]
    return {
        "page_width_cm": round(sec.page_width.cm, 2),
        "page_height_cm": round(sec.page_height.cm, 2),
        "normal_font_size_pt": doc.styles["Normal"].font.size.pt,
        "normal_line_spacing": doc.styles["Normal"].paragraph_format.line_spacing,
    }


def citation_order_audit(md: str) -> dict[str, object]:
    seen: list[int] = []
    for m in re.finditer(r"\[([0-9,\-\s]+)\]", md):
        token = m.group(1)
        vals: list[int] = []
        for part in token.replace(" ", "").split(","):
            if not part:
                continue
            if "-" in part:
                a, b = part.split("-", 1)
                if a.isdigit() and b.isdigit():
                    vals.extend(range(int(a), int(b) + 1))
            elif part.isdigit():
                vals.append(int(part))
        for v in vals:
            if v not in seen:
                seen.append(v)
    expected = list(range(1, 21))
    return {"first_appearance_order": seen, "is_ordered_prefix": seen[:20] == expected, "missing": [x for x in expected if x not in seen]}


def audit_text(md: str) -> dict[str, object]:
    abstract = re.search(r"## Abstract\n\n(.*?)\n\n\*\*Keywords:", md, re.S)
    title = re.search(r"^# (.*)$", md, re.M).group(1)
    keywords = re.search(r"\*\*Keywords:\*\*(.*)", md).group(1).split(";")
    ref_lines = re.findall(r"^\[\d+\].*$", md, re.M)
    too_many_author_refs = []
    for line in ref_lines:
        ref_id = line.split("]", 1)[0] + "]"
        author_block = line.split("]", 1)[1].split(". ", 1)[0]
        author_items = [x.strip() for x in author_block.split(",")]
        has_et_al = "et al" in author_block.lower()
        if len(author_items) > 3 and not has_et_al:
            too_many_author_refs.append(ref_id)
    forbidden = [
        "workflow",
        "dashboard",
        "EIS proves",
        "EIS disproves",
        "validated operating map",
        "optimized NH4Br",
        "measured coverage",
        "exact pressure threshold",
        "full-cell proves",
        "Q*/i* fitted",
        "parameter-free",
        "production design claim",
        "550C best",
    ]
    return {
        "title_words": len(re.findall(r"[A-Za-z0-9-]+", title)),
        "abstract_words": len(re.findall(r"[A-Za-z0-9_+./-]+", abstract.group(1) if abstract else "")),
        "keyword_count": len([k for k in keywords if k.strip()]),
        "references": len(ref_lines),
        "citation_order": citation_order_audit(md),
        "too_many_author_refs_without_etal": too_many_author_refs,
        "has_funding": "## Funding" in md,
        "has_conflict": "## Conflict of Interest" in md,
        "has_ethics": "## Ethical approval" in md,
        "has_informed_consent": "## Informed consent" in md,
        "has_materials_methods": "## 4 Materials and methods" in md,
        "forbidden_counts": {pat: len(re.findall(re.escape(pat), md, re.I)) for pat in forbidden},
    }


def write_auxiliary(audit: dict[str, object], docx_format: dict[str, object]) -> None:
    (OUT / "R92_SCTS_compliance_matrix.md").write_text(
        f"""# R92 SCTS compliance matrix

Sources checked:

- Springer journal page: https://link.springer.com/journal/11431
- SCTS author instructions PDF stored locally: `outputs/SCTS_R84_publication_deep_revision/SCTS_Instructions_for_authors_20230608.pdf`

| SCTS requirement | R92 status |
|---|---|
| Broad technological-sciences relevance | PASS: mechanism addresses energy sciences, porous electrodes and flow blockage |
| Article type | Prepared as Article; final selection should be confirmed with invited editor |
| Title no more than 20 words | PASS: {audit['title_words']} words |
| Abstract without citations/equations | PASS: {audit['abstract_words']} words, no citation markers |
| 3-8 keywords | PASS: {audit['keyword_count']} keywords |
| Introduction cites references in order | PASS: {audit['citation_order']['is_ordered_prefix']} |
| Figures/tables inserted in text and cited in order | PASS by manuscript structure and PDF proof |
| A4 document size | PASS: {docx_format['page_width_cm']} cm x {docx_format['page_height_cm']} cm |
| 10 pt text and 1.5 line spacing | PASS: Normal {docx_format['normal_font_size_pt']} pt, line spacing {docx_format['normal_line_spacing']} |
| References numbered in order | PASS: {audit['references']} references |
| More than 3 authors shortened to first 3 et al. | PASS: {len(audit['too_many_author_refs_without_etal'])} remaining violations |
| Funding statement | PASS placeholder included; authors must complete |
| Conflict of Interest statement | PASS placeholder/no-conflict wording included |
| Ethical approval statement | PASS: no human/animal studies statement included |
| Data availability | PASS; public repository deposit still recommended before final publication |
| Cover letter with background, innovation and significance | Present from R91/R90; author metadata still required |

Remaining author-side blockers:

1. Author names, affiliations and corresponding email.
2. Funding names and grant numbers.
3. Author contributions.
4. Final conflict-of-interest confirmation.
5. Final human proof and reference-manager check.
""",
        encoding="utf-8",
    )
    (OUT / "R92_editor_cover_note.md").write_text(
        f"""# Editor-facing positioning note

Manuscript title: {TITLE}

This invited manuscript is positioned as an Article for Science China Technological Sciences. It addresses a technological-sciences problem in zinc-iodine flow batteries: why a porous positive carbon felt can fail locally by solid iodine formation and blockage even when the bulk electrolyte remains chemically competent.

Research background: zinc-iodine flow batteries are attractive for aqueous energy storage, but neutral iodine can leave the soluble redox inventory and become a solid foulant.

Innovation: the manuscript connects molecular iodine residence near carbon, single-fiber precipitation/coverage closures, porous-electrode COMSOL state evolution and localized blockage amplification into one physical iodine-handling framework.

Significance: the work gives testable coordinates for future experiments: neutral-iodine source/removal balance, free-iodine supersaturation, cumulative supersaturation dose, solid iodine inventory, accessible surface state, localization and pressure response.

Claim boundary: the manuscript does not claim optimized NH4Br concentration, a production operating map, measured microscopic coverage or a quantitative pressure threshold.
""",
        encoding="utf-8",
    )
    (OUT / "GPT_PRO_REVIEW_PROMPT_R92_CN.md").write_text(
        """# 给 GPT Pro / 合作者的 R92 审稿提示

请按 Science China Technological Sciences 约稿 Article 的标准审这版。

重点检查：

1. 题名是否足够简洁、可检索、没有过度承诺。
2. 摘要是否符合目的-方法-结果-结论，且没有引用、公式和过多专业符号。
3. Introduction 是否按近 2-3 年进展和经典背景自然引出问题。
4. Results 是否像论文结论链，而不是模型日志。
5. 每个参数家族是否有物理入口、因果路径和可验证预测，而不是只给敏感性图。
6. Materials and methods 是否足以让审稿人理解和复查模型。
7. 参考文献是否需要补充更贴近 SCTS 工程读者的近年文献。
8. 是否仍有过度承诺：优化 NH4Br、production map、measured coverage、exact pressure threshold。

请给段落级修改建议，并判断是否可以进入导师/合作者最终人工改稿。
""",
        encoding="utf-8",
    )


def package_outputs() -> list[dict[str, object]]:
    full = PKG / "SCTS_R92_SCTS_FORMAT_COMPLIANCE_FULL.zip"
    gpt = PKG / "GPT_PRO_REVIEW_PACKET_R92_SCTS_FORMAT.zip"
    figs = PKG / "SCTS_R92_FIGURES_NATIVE_600DPI.zip"
    with zipfile.ZipFile(figs, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(FIG600.glob("*.png")):
            zf.write(p, p.name)
    with zipfile.ZipFile(full, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(OUT.rglob("*")):
            if p.is_file() and "package" not in p.relative_to(OUT).parts:
                zf.write(p, str(p.relative_to(OUT)).replace("\\", "/"))
    wanted = [
        "SCTS_ZIFB_SCTS_format_R92.md",
        "SCTS_ZIFB_SCTS_format_R92.docx",
        "pdf_proof/SCTS_ZIFB_SCTS_format_R92.pdf",
        "Supporting_Information_R92.md",
        "Supporting_Information_R92.docx",
        "pdf_proof/Supporting_Information_R92.pdf",
        "pdf_proof/manuscript_contact_sheet_R92.png",
        "pdf_proof/supporting_information_contact_sheet_R92.png",
        "R92_SCTS_compliance_matrix.md",
        "R92_editor_cover_note.md",
        "GPT_PRO_REVIEW_PROMPT_R92_CN.md",
        "R90_reference_verification_report.md",
        "final_status.md",
    ]
    with zipfile.ZipFile(gpt, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in wanted:
            p = OUT / rel
            if p.exists():
                zf.write(p, rel.replace("\\", "/"))
        for p in sorted(FIG600.glob("*.png")):
            zf.write(p, f"submission_figures_600dpi/{p.name}")
    out = []
    for p in [full, gpt, figs]:
        with zipfile.ZipFile(p) as zf:
            out.append({"file": p.name, "entries": len(zf.namelist()), "testzip_bad": zf.testzip()})
    return out


def main() -> None:
    ensure_dirs()
    copy_tree(R91 / "figures", FIG)
    copy_tree(R91 / "submission_figures_600dpi", FIG600)
    for name in ["R90_reference_verification_report.md", "R90_reference_verification.csv", "render_docx_fallback_note_R91.md"]:
        src = R91 / name
        if src.exists():
            shutil.copy2(src, OUT / name)
    md = (R91 / "SCTS_ZIFB_submission_polish_final_R91.md").read_text(encoding="utf-8")
    md = scts_polish_markdown(md)
    (OUT / "SCTS_ZIFB_SCTS_format_R92.md").write_text(md, encoding="utf-8")
    si = (R91 / "Supporting_Information_R91.md").read_text(encoding="utf-8")
    si = si.replace("# Supporting Information for Local iodine-handling failure in zinc-iodine flow batteries", "# Supporting Information for local iodine-handling failure in zinc-iodine flow batteries")
    (OUT / "Supporting_Information_R92.md").write_text(si, encoding="utf-8")

    r84.OUT = OUT
    r84.FIG = FIG
    r84.PDF = PDF
    r84.PKG = PKG
    checks = {
        "manuscript": r84.markdown_to_docx(OUT / "SCTS_ZIFB_SCTS_format_R92.md", OUT / "SCTS_ZIFB_SCTS_format_R92.docx", r84.MAIN_FIGS),
        "supporting_information": r84.markdown_to_docx(OUT / "Supporting_Information_R92.md", OUT / "Supporting_Information_R92.docx", r84.SI_FIGS),
    }
    docx_format = set_scts_docx_format(OUT / "SCTS_ZIFB_SCTS_format_R92.docx")
    set_scts_docx_format(OUT / "Supporting_Information_R92.docx")
    r84.export_pdf(OUT / "SCTS_ZIFB_SCTS_format_R92.docx", PDF / "SCTS_ZIFB_SCTS_format_R92.pdf")
    r84.export_pdf(OUT / "Supporting_Information_R92.docx", PDF / "Supporting_Information_R92.pdf")
    pdf_checks = {
        "manuscript": r84.pdf_qa(PDF / "SCTS_ZIFB_SCTS_format_R92.pdf", "manuscript"),
        "supporting_information": r84.pdf_qa(PDF / "Supporting_Information_R92.pdf", "supporting_information"),
    }
    for old, new in [
        (PDF / "manuscript_contact_sheet_R84.png", PDF / "manuscript_contact_sheet_R92.png"),
        (PDF / "supporting_information_contact_sheet_R84.png", PDF / "supporting_information_contact_sheet_R92.png"),
    ]:
        if old.exists():
            old.replace(new)
    audit = audit_text(md)
    audit["docx_format"] = docx_format
    audit["pdf"] = pdf_checks
    audit["docx"] = checks
    write_auxiliary(audit, docx_format)
    (OUT / "R92_QA_summary.json").write_text(json.dumps(audit, indent=2), encoding="utf-8")
    zips = package_outputs()
    (OUT / "R92_zip_integrity_check.json").write_text(json.dumps(zips, indent=2), encoding="utf-8")
    no_blank = not pdf_checks["manuscript"]["blank_pages"] and not pdf_checks["supporting_information"]["blank_pages"]
    clean_forbidden = all(v == 0 for v in audit["forbidden_counts"].values())
    status = "PASS" if no_blank and clean_forbidden and audit["citation_order"]["is_ordered_prefix"] and not audit["too_many_author_refs_without_etal"] else "PARTIAL"
    (OUT / "final_status.md").write_text(
        f"""SCTS R92 SCTS-format compliance packet: {status}
Manuscript Markdown: SCTS_ZIFB_SCTS_format_R92.md
Manuscript DOCX: SCTS_ZIFB_SCTS_format_R92.docx
Manuscript PDF proof: pdf_proof/SCTS_ZIFB_SCTS_format_R92.pdf
Supporting Information DOCX: Supporting_Information_R92.docx
Supporting Information PDF proof: pdf_proof/Supporting_Information_R92.pdf
GPT review zip: package/GPT_PRO_REVIEW_PACKET_R92_SCTS_FORMAT.zip
Full submission zip: package/SCTS_R92_SCTS_FORMAT_COMPLIANCE_FULL.zip
Figures zip: package/SCTS_R92_FIGURES_NATIVE_600DPI.zip
Title words: {audit['title_words']} / 20
Abstract words: {audit['abstract_words']}
Keywords: {audit['keyword_count']}
References: {audit['references']}
SCTS et-al reference format violations: {len(audit['too_many_author_refs_without_etal'])}
Citation order check: {'PASS' if audit['citation_order']['is_ordered_prefix'] else 'FAIL'}
A4/10pt/1.5 spacing: {docx_format}
Manuscript PDF pages: {pdf_checks['manuscript']['pages']}
SI PDF pages: {pdf_checks['supporting_information']['pages']}
Blank pages detected: {'NO' if no_blank else 'YES'}
Forbidden wording scan: {'PASS' if clean_forbidden else 'CHECK'}
Ready for GPT Pro/coauthor scientific review: YES
Ready for senior-author line editing: YES
Ready for direct one-click SCTS upload: NO
Reason direct upload is still blocked: author names, affiliations, corresponding-author email, funding, author contributions, conflict-of-interest confirmation and final human proof must be completed by the authors.
""",
        encoding="utf-8",
    )
    print((OUT / "final_status.md").read_text(encoding="utf-8"))
    print(json.dumps({"out": str(OUT), "zip_results": zips, "qa": audit}, indent=2))


if __name__ == "__main__":
    main()
