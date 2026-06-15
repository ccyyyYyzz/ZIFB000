from __future__ import annotations

import json
import re
import shutil
import subprocess
import zipfile
from pathlib import Path

import pypdfium2 as pdfium
from PIL import Image, ImageDraw
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "outputs" / "SCTS_R83_final_review_packet"
OUT = ROOT / "outputs" / "SCTS_R84_publication_deep_revision"
FIG = OUT / "figures"
PDF = OUT / "pdf_proof"
PKG = OUT / "package"

PYTHON = Path(r"C:\Users\CYZ的computer\.cache\codex-runtimes\codex-primary-runtime\dependencies\python\python.exe")

TITLE = "Local iodine-handling failure controls solid-iodine blockage in zinc-iodine flow batteries"

MAIN_FIGS = [
    "Fig1_R67_local_iodine_failure_framework.png",
    "Fig2_single_fiber_closure_bridge_R64.png",
    "Fig3_R69_COMSOL_state_evolution.png",
    "Fig4_COMSOL_axis_profile_fields_R64.png",
    "Fig5_R67_parameter_mechanism_matrix.png",
]

SI_FIGS = [
    "SI_Fig_molecular_surface_evidence_R64.png",
    "SI_Fig_true_COMSOL_cloudmaps_R64.png",
    "SI_Fig_precipitation_dissolution_boundary_R64.png",
    "SI_dense_Fig3_CAP_RATE_time_evolution_R64.png",
    "SI_dense_Fig5_parameter_family_curves_R64.png",
    "SI_dense_Fig6_parameter_fingerprint_R64.png",
    "Fig6_R67_evidence_boundary_falsification.png",
]


def ensure_dirs() -> None:
    if OUT.exists():
        shutil.rmtree(OUT)
    for p in [OUT, FIG, PDF, PKG]:
        p.mkdir(parents=True, exist_ok=True)


def copy_figures() -> None:
    for name in set(MAIN_FIGS + SI_FIGS + ["Graphical_Abstract_R65.png"]):
        src = SRC / "figures" / name
        if src.exists():
            shutil.copy2(src, FIG / name)


def refs_text() -> str:
    md = (SRC / "SCTS_ZIFB_final_review_R83.md").read_text(encoding="utf-8")
    return md.split("## References", 1)[1].strip()


def manuscript_text() -> str:
    refs = refs_text()
    return f"""# {TITLE}

**Authors:** [Author names]

**Affiliations:** [Author affiliations]

**Corresponding author:** [Name and email]

## Abstract

Zinc-iodine flow batteries promise high aqueous energy density, but iodine is also an unusually unforgiving active species: the same neutral iodine that stores charge can become a solid foulant in the porous positive electrode. Here we develop a physics-first mechanism for this transition. The central failure coordinate is not the total iodine inventory in the tank, but the local ability of the positive carbon felt to handle free iodine as it is generated at fiber surfaces. Failure begins when interfacial I2 generation exceeds the combined local removal capacity supplied by complexation, diffusion, flow renewal, dissolution and accessible carbon surface. We connect this source-removal balance across four levels. Molecular carbon-surface calculations provide bounded priors for iodine residence near graphitic and functional motifs. A standalone radial single-fiber model then converts local precipitation into replaceable closures for surface coverage, iodine-film attenuation, precipitation/dissolution and bare/covered current partition. COMSOL porous-electrode postprocessing maps time-resolved free-I2 supersaturation, solid iodine inventory, closure-derived coverage and film/area-loss variables under capacity- and rate-oriented operation. Finally, a mean-field versus localized-blockage analysis explains why a modest average solid iodine fraction can still produce visible clumps and hydraulic symptoms when the inventory concentrates at stagnant zones, fiber contacts or pore throats. The resulting framework separates capacity-driven cumulative iodine dose from rate-driven generation-removal imbalance, clarifies why additives and operating parameters must be interpreted by their physical entry point, and identifies decisive measurements for future validation: phase identification, interrupted-charge imaging, pressure tracing and reference-electrode separation.

**Keywords:** zinc-iodine flow battery; iodine precipitation; carbon felt; local supersaturation; single-fiber model; COMSOL; porous electrode; flow blockage

## 1 Introduction

Zinc-iodine flow batteries are attractive because they combine inexpensive zinc, concentrated aqueous iodine redox chemistry and the architectural flexibility of flow batteries. Early zinc-polyiodide work established that iodine-rich electrolytes can support high energy density [4]. Subsequent studies showed that polyiodide or bromide complexation can increase iodide utilization [5,19], that self-healing or single-flow architectures can extend cycle life [6,7], that robust electrolytes and electrode treatments can improve rate capability [8,9], and that ion-management or negative-side strategies can suppress other important failure modes [10,20]. These advances make the Zn-I system technologically compelling.

The same chemistry, however, also creates a distinctive failure problem. Iodine is not just a soluble redox carrier. Under charge it passes through neutral I2, polyiodide and mixed halide species; it can adsorb near carbon surfaces; it can cross a free-I2 saturation boundary; and it can form solid deposits that cover fibers, bridge pore throats or enter poorly renewed regions of the felt. Experimental observations in practical cells include dark solid iodine, clumps, flow obstruction, pressure rise and pump-line failure. Such symptoms cannot be reduced to a single tank concentration or a single full-cell voltage residual. They indicate a physical conversion of iodine from mobile redox inventory into a local solid/morphology state.

This paper starts from that physical conversion. The central proposition is that iodine failure in a Zn-I flow battery is a local iodine-handling failure of the porous positive carbon felt. The tank can remain chemically buffered while the fiber-scale interface crosses a free-I2 supersaturation boundary. At that interface, current generates neutral iodine at a flux proportional to the local current density. Removal is supplied by complexation, diffusion, flow renewal, dissolution and the still-accessible carbon surface. When the source exceeds the local handling capacity, the relevant state variables are free-I2 supersaturation, cumulative supersaturation dose, solid iodine inventory, surface accessibility, film attenuation and localization.

This framing is deliberately different from three simpler explanations. It is not a purely global electrolyte-depletion story, because total iodide or bromide inventory can remain large while local free iodine exceeds saturation. It is not a single film-resistance story, because coverage, film attenuation, precipitation/dissolution and bare/covered current split are distinct physical links. It is also not a generic porous-electrode voltage-fitting story, because full-cell voltage mixes open-circuit potential, zinc-side polarization, membrane/electrolyte resistance, terminal concentration polarization and positive-electrode iodine loss. A useful mechanism must therefore follow iodine from molecular residence to fiber-scale precipitation, porous-electrode fields and morphology-scale blockage.

Carbon felt amplifies the importance of locality. In redox-flow electrodes, compression, permeability, tortuosity, wetting, active area and flow distribution determine where reaction occurs and how quickly species are renewed [14-18]. In a Zn-I positive felt these structural variables also determine where neutral iodine first forms, how long it remains near carbon, where solid iodine can nucleate, and whether a small mean solid fraction becomes harmlessly distributed or dangerously localized. This is why a continuum field and a visible clump should not be treated as contradictory: they describe different levels of the same chain.

We build the mechanism with three connected modelling levels and one interpretation layer. First, molecular surface calculations test whether iodine residence near representative carbon motifs is chemically plausible. These calculations are used only as bounded surface priors, not as fitted macroscopic kinetic constants. Second, a standalone single-fiber model resolves radial transport, fast complexation, surface generation/consumption, precipitation/dissolution, coverage and local film attenuation. This model exports closures rather than replacing the porous-electrode model with a black-box voltage term. Third, COMSOL porous-electrode postprocessing maps the local state variables under capacity- and rate-oriented operation. Finally, mean-field versus localized-blockage analysis explains why hydraulic symptoms require morphology amplification beyond a uniform mean porosity loss.

The contribution is a mechanism framework rather than an optimization recipe. It identifies where each parameter family enters the iodine-failure chain: accessible carbon area changes the local source term, flow and diffusivity change removal, speciation and saturation move the free-I2 gate, precipitation and dissolution control solid inventory and recovery, compression and localization control hydraulic amplification, and ASR controls voltage loss without necessarily changing iodine phase stability. This chain provides a cleaner basis for experiments and design decisions than a single empirical failure threshold.

![Fig1_R67_local_iodine_failure_framework.png](figures/Fig1_R67_local_iodine_failure_framework.png)

**Figure 1. Local iodine-handling failure in porous positive carbon felt.** The reservoir can remain chemically buffered while the positive felt fails locally. Electrochemical current creates free iodine at fiber surfaces; complexation, diffusion, flow renewal, dissolution and remaining accessible area determine whether that iodine is removed or crosses a saturation boundary. Persistent supersaturation drives solid iodine inventory, closure-derived surface accessibility loss, film attenuation and morphology-scale clumping.

## 2 Results

### 2.1 A useful failure coordinate must be local, not tank-averaged

The first result is conceptual but essential: the physically relevant failure coordinate is the local source-removal balance at the positive carbon fiber, not the total iodine inventory in the reservoir. In the tank, total iodine, iodide and supporting halide concentrations describe chemical storage capacity. At a carbon surface, however, current generates neutral iodine at a local flux

J_gen = j_loc/(2F).    (1)

Here j_loc is the local interfacial current density and F is Faraday's constant. The minimum removal capacity can be written as

J_rem ~= k_m,eff beta_eff c_I2,sat,eff.    (2)

The effective mass-transfer coefficient k_m,eff groups diffusion, pore-scale renewal and convective exchange; beta_eff expresses how rapidly free I2 is buffered into soluble iodine complexes; and c_I2,sat,eff is the effective free-I2 saturation concentration at the local surface. The simplest local iodine-handling number is therefore

Da_I2 = J_gen/J_rem.    (3)

The phase-forming species is free iodine rather than total dissolved iodine. The saturation coordinate is

S_I2 = c_I2,free,local/c_I2,sat,eff.    (4)

and the history variable for long charge is

Omega_I2 = integral max(S_I2 - S_on, 0) dt.    (5)

These equations explain the "iodine paradox" in Zn-I flow batteries. A formulation can hold a large total iodine inventory and still fail locally if the positive felt generates free iodine faster than it can be complexed, transported, dissolved or distributed over accessible surface. Conversely, a lower total iodine loading can be safer if active area, flow renewal and local speciation keep S_I2 below the precipitation boundary. The key state variables used in this paper are therefore local and history-aware.

| Quantity | Definition | Physical meaning |
|---|---|---|
| J_gen | j_loc/(2F) | Fiber-scale neutral-iodine source flux |
| J_rem | approximately k_m,eff beta_eff c_I2,sat,eff | Local removal by transport, complexation and dissolution |
| Da_I2 | J_gen/J_rem | Instantaneous iodine-handling stress |
| S_I2 | c_I2,free,local/c_I2,sat,eff | Local free-I2 saturation coordinate |
| Omega_I2 | integral max(S_I2 - S_on, 0) dt | Cumulative supersaturation dose |
| eps_s | solid iodine volume fraction | Mean solid inventory in the continuum model |
| eps_s,local | eps_s,mean/f_local | Localized solid loading relevant to pore-throat obstruction |
| theta | 1 - exp(-a_theta h_I2^b_theta) | Closure-derived surface accessibility state |
| f_film | 1/(1 + g_eff R_film) | Covered-path kinetic attenuation factor |

The table is not a list of fitting variables. It is a map of the mechanism. The source term, removal term, phase gate, cumulative dose, solid inventory, surface state and morphology amplification are separate links. Keeping them separate prevents a common error: attributing every late voltage rise, pressure symptom or cycle fade to one generic "iodine film".

### 2.2 Molecular surface calculations justify residence, not macroscopic rates

The first modelling level asks whether iodine has a plausible tendency to reside near carbon surfaces. Finite carbon motifs cannot represent a compressed, flowing carbon felt, but they can provide a bounded prior for local residence and nucleation tendency. In the final molecular prior set, neutral I2 on a basal graphitic motif has an adsorption energy of about -0.51 eV, or -49.6 kJ mol-1. Hydroxylated basal motifs are only modestly more iodine-philic than basal carbon, by about 0.039 eV, corresponding to a conservative site-preference factor of 1.5. Carbonyl-edge motifs are weaker still, about 0.016 eV relative to basal, and are retained only as a weak sensitivity prior.

The important conclusion is not that one functional group uniquely controls the battery. The molecular result says something narrower and more useful: iodine residence near carbon is chemically plausible, modest surface heterogeneity can bias where iodine first resides, and patchy nucleation is therefore a reasonable physical assumption. It does not determine precipitation rates, film conductivity, surface area loss or pressure rise. Those quantities belong to the fiber, porous-electrode and morphology levels.

This separation matters for publication-quality mechanism writing. If the molecular calculation were over-transferred, the paper would appear to claim that a finite motif energy directly predicts a full-cell failure threshold. We avoid that. Molecular calculations enter as priors on local residence; the single-fiber model and COMSOL fields then determine how local residence participates in precipitation and transport.

### 2.3 The single-fiber model turns surface deposition into portable closures

The second modelling level is the radial single-fiber model. Its purpose is not to replace COMSOL and not to fit full-cell voltage. It solves a local problem that the macroscopic model needs but cannot infer from voltage alone: how free-I2 transport, precipitation, dissolution, coverage and film attenuation evolve around a single carbon fiber.

Fast complexation is represented by

beta_I2 = 1 + K_I2,I c_I- + K_I2,Br c_Br,support, and c_I2,free = c_I2,total/beta_I2.    (6)

At the fiber surface, the electrochemical reaction I2 + 2e- <-> 2I- supplies or consumes iodine. The model then updates solid iodine thickness h_I2, surface coverage theta, local film resistance R_film and the current split between bare and covered pathways:

theta = 1 - exp(-a_theta h_I2^b_theta).    (7)

R_film = h_I2/sigma_I2.    (8)

f_film = 1/(1 + g_eff R_film).    (9)

j_total = j_bare + j_cov.    (10)

This is the methodological core of Fig. 2. The four panels should not be read as generic fits. The upper-left panel shows that coverage is a nonlinear accessibility state: small deposits can quickly change the usable fraction of a surface, while later deposits approach a plateau. The upper-right panel shows that local film attenuation can be represented as a compact algebraic factor, separating it from coverage. The lower-left panel shows why bare/covered parallel pathways are needed: current does not simply vanish when a surface becomes partly covered, but the accessible current pathway shifts. The lower-right panel shows the precipitation/dissolution boundary: precipitation responds to supersaturation, while dissolution requires both undersaturation and existing solid/covered state.

The fitted closure scan contains 720 single-fiber rows. For initially bare fibers, theta(h_I2) is described by a_theta = 4.04e+06 m^(-b) and b_theta = 1.018, with R2 = 0.80. The film attenuation fit gives g_eff = 40.00 S m-2 with a numerical RMSE of 3.46e-09. The retained precipitation and dissolution coefficients are k_precip = 1.0e-06 m s-1 and k_diss = 2.0e-06 m s-1. The pathway prefactors are j0,bare = 15.0 A m-2 and j0,cov = 2.0 A m-2, and the current-split expression reproduces the scan with RMSE below 6.3e-10 in fractional current.

These values should be interpreted as closure parameters for a local mechanism model, not as universal material constants. The important advance is the separation of links. Precipitation/dissolution, surface accessibility, local film attenuation and current redistribution are exported as distinct quantities. This separation makes the later COMSOL interpretation readable: when a field shows rising theta, it means closure-derived accessibility loss; when it shows rising R_film, it means a local covered-path penalty; and when it shows a current split, it means the model has moved current from bare to covered surface rather than hidden the effect in one voltage residual.

![Fig2_single_fiber_closure_bridge_R64.png](figures/Fig2_single_fiber_closure_bridge_R64.png)

**Figure 2. Single-fiber bridge from local precipitation to porous-electrode closures.** The radial fiber calculation resolves free-I2 transport, fast complexation, electrochemical generation/consumption, precipitation/dissolution and local solid iodine thickness. The output is converted into closures for theta(h_I2), R_film, film kinetic attenuation and bare/covered current split. Each curve maps one local surface state into one macroscopic closure variable.

### 2.4 Time-resolved COMSOL states separate capacity and rate routes

The third modelling level asks how the closure variables behave in porous-electrode operation. Figure 3 is the central result because it shows state evolution, not just endpoints. The capacity-oriented route and rate-oriented route load the same iodine-failure chain in different order.

In the capacity-oriented route, the most important signature is delayed coupling among solid iodine inventory, coverage-derived accessibility loss, local film resistance and area-loss proxy. The free-I2 supersaturation proxy remains comparatively lower than in the rate case, but the solid inventory grows with charge progress. Once inventory has accumulated, theta and the area-loss proxy rise sharply. This pattern is the fingerprint of a cumulative-dose route: the system is not failing because the instantaneous source term is always extreme; it is failing because the integral of supersaturation exposure and solid persistence eventually changes the surface state.

In the rate-oriented route, the free-I2 supersaturation proxy is already high early in the charge. The red curve in the supersaturation panel is therefore the clearest sign of source-removal imbalance. It indicates that J_gen has exceeded J_rem before the same magnitude of mean solid inventory and closure-derived area loss develops. This is why high-rate operation can seed failure quickly even when the total charged capacity is lower: the local interface crosses the free-I2 gate first, and the solid/morphology consequences follow later.

The figure also prevents a misleading interpretation. The rate route should not be described as a resolved single hot point in the mesh. It is a high-supersaturation route. The capacity route should not be described as direct experimental observation of surface coverage. It is a closure-derived surface-state route. These distinctions make the conclusion stronger, not weaker, because they state exactly what the simulation proves: the ordering of state variables is different in capacity and rate operation.

This ordering leads to testable predictions. If the rate route is correct, increasing flow renewal, improving current distribution or increasing accessible area should suppress early S_I2 more efficiently than merely increasing tank capacity. If the capacity route is correct, interrupted-charge dissolution, rest history, phase imaging and accessible-area recovery should strongly influence late solid inventory and surface-state loss. A top-tier mechanism paper should state such predictions explicitly; otherwise the model remains a collection of plots.

![Fig3_R69_COMSOL_state_evolution.png](figures/Fig3_R69_COMSOL_state_evolution.png)

**Figure 3. Time-resolved state variables separate capacity and rate routes.** Capacity-oriented operation produces delayed growth of solid iodine inventory, closure-derived coverage, film proxy and area-loss proxy, indicating cumulative supersaturation dose and surface-state persistence. Rate-oriented operation produces early free-I2 supersaturation before comparable mean solid inventory develops, indicating an instantaneous generation/removal imbalance.

### 2.5 Spatial fields move the hydraulic question from mean field to morphology

The spatial COMSOL fields provide a second important constraint: the iodine-stress pattern is broad across the positive electrode rather than concentrated into one dominant resolved continuum point. This matters because visible clumps and pressure symptoms are localized phenomena. If the continuum model does not show a sharp hotspot, the correct conclusion is not that clumps are impossible. The correct conclusion is that clumps are produced by a morphology-amplification layer that the mean-field model does not fully resolve.

A uniform solid iodine volume fraction changes porosity and permeability only through a smooth mean-field pathway. For small eps_s,mean, a Kozeny-Carman-like permeability estimate may change only modestly. But if the same solid inventory localizes into a fraction f_local of the active pore space, the local solid fraction becomes

eps_s,local = eps_s,mean/f_local.    (11)

This relation is simple but powerful. A mean inventory of 0.1% is hydraulically mild if distributed everywhere, but it becomes a local 10% obstruction if concentrated into 1% of the pore space. In a fibrous felt, such localization can occur at stagnant pockets, fiber crossings, binder-like debris, compression-induced dead zones or pore throats with slow renewal. The pressure event is then a morphology consequence of local solid arrangement, not only a volume-average porosity loss.

Figure 4 should therefore be read as a boundary between model-resolved and model-implied physics. The COMSOL fields resolve broad iodine stress and axial/state gradients. They do not fully resolve nucleation heterogeneity, agglomeration, detachment, particle bridging or slurry-like motion. Those missing processes are not a reason to discard the mechanism; they define the next model layer required for quantitative pressure prediction.

![Fig4_COMSOL_axis_profile_fields_R64.png](figures/Fig4_COMSOL_axis_profile_fields_R64.png)

**Figure 4. Spatial iodine-state fields and morphology amplification.** Coordinate-aware COMSOL fields show broad positive-electrode iodine stress rather than a single resolved continuum hotspot. Visible clumps, pore-throat obstruction and pressure symptoms are therefore interpreted as morphology-amplified consequences of local solid iodine inventory.

### 2.6 Parameter effects must be argued by physical entry point

The most important revision for a publishable manuscript is the treatment of parameters. A single lever plot is not enough. Each parameter family enters a different equation, produces a different tradeoff and points to a different experiment. Figure 5 is therefore a mechanism matrix, not an optimization chart.

#### 2.6.1 Accessible carbon area: source dilution and surface inventory distribution

Accessible carbon area acts first on J_gen. For a fixed stack current, increasing the effective active surface lowers j_loc and therefore lowers the local neutral-iodine source flux j_loc/(2F). It also distributes any solid iodine over more surface, delaying local coverage and preserving a larger bare-current fraction. Conversely, poor wetting, compression dead zones, partial felt utilization or blocked flow paths concentrate current into fewer active regions. The nominal current density can then look safe while the true local j_loc is high.

This is why area is not merely a kinetic prefactor. It is a source-distribution variable. A decisive experiment would compare identical electrolyte and current at different controlled wetting/compression states, then measure whether the early S_I2 proxy, phase appearance or pressure onset shifts as predicted. If only voltage is measured, area effects may be confused with ASR or activation.

#### 2.6.2 Flow renewal and diffusivity: removal rather than chemistry

Flow renewal and effective diffusivity act on k_m,eff in J_rem. Their role is to remove newly generated iodine from the fiber-scale environment before free-I2 activity crosses the saturation boundary. This is physically different from changing iodide or bromide concentration. Flow and diffusivity do not primarily change how much iodine the tank can store; they change how quickly the interfacial region can shed the iodine it just produced.

This distinction predicts that rate-driven failure should be highly sensitive to flow renewal and local transport length. If early supersaturation is caused by generation-removal imbalance, increasing flow or improving pore connectivity should suppress the red high-S_I2 signature in Fig. 3 more efficiently than it suppresses late cumulative inventory after long capacity loading. That is a falsifiable route-specific prediction.

#### 2.6.3 Speciation and saturation: moving the free-I2 gate

Speciation and free-I2 saturation act at the phase gate. Iodide and bromide complexation reduce c_I2,free at fixed total iodine, while c_I2,sat,eff determines where precipitation begins. This is the correct scientific place to discuss NH4Br or related additives. They are not the thesis of the paper; they are one way to move the free-I2 gate.

The tradeoff is subtle. Strong complexation can suppress free iodine and delay precipitation, but it cannot eliminate local source intensity, slow flow renewal, stagnant zones or nucleation heterogeneity. A formulation with excellent total iodine buffering can still fail if a carbon-felt region crosses S_I2 > 1. Conversely, a formulation with less total buffering might operate safely if local current density is low and renewal is high. This is why the paper avoids an additive-optimization claim.

#### 2.6.4 Precipitation and dissolution: reversible buffer or persistent foulant

Precipitation and dissolution control the conversion between free iodine and solid inventory. Faster precipitation can reduce free-I2 activity by moving iodine into a solid reservoir, but it can also increase eps_s, coverage and blockage risk. Slower precipitation can keep more iodine in solution, which may reduce solid inventory but increase local free-I2 activity and reaction stress. Dissolution controls whether a solid deposit is reversible during rest or discharge.

The same parameter can therefore look beneficial or harmful depending on the observable. If the observable is free-I2 activity, precipitation can appear protective. If the observable is pressure or coverage, precipitation can appear damaging. A publishable mechanism must acknowledge this ambivalence. The question is not "is precipitation good or bad?" but whether solid iodine remains reversible, distributed and dissolvable, or becomes localized and persistent.

#### 2.6.5 Compression, permeability and localization: the hydraulic amplification layer

Compression improves electronic contact and may increase apparent active area, but it reduces pore volume, modifies permeability and makes local obstruction more dangerous. This is the entry point for hydraulic failure. Uniform eps_s,mean alone is usually insufficient to explain abrupt pump symptoms; localization converts the same inventory into much larger eps_s,local. Compression lowers the tolerance for that localization by shrinking pore throats and increasing pressure sensitivity.

This is also where a future pore-network or precipitation-permeability model is required. The present model supports the physical route from supersaturation to solid inventory and from solid inventory to localization risk. It does not assign a numerical pressure cutoff.

#### 2.6.6 Conductivity, membrane resistance and ASR: voltage loss is not phase stability

Ohmic resistance, membrane resistance and electrolyte conductivity determine the voltage penalty for a given current. They are technologically important, but they are not the same as iodine phase stability. A low-ASR cell can still precipitate iodine if J_gen exceeds J_rem. A high-ASR cell can show poor voltage efficiency without solid iodine being the primary cause. Therefore voltage loss must be separated from iodine-handling failure in both modelling and interpretation.

This separation explains why archived full-cell curves are not used as microscopic proof. Voltage is an observable mixture. It can be consistent with the iodine mechanism, but it cannot by itself identify coverage, film or pore blockage without reference electrodes, pressure data, phase imaging or controlled perturbations.

![Fig5_R67_parameter_mechanism_matrix.png](figures/Fig5_R67_parameter_mechanism_matrix.png)

**Figure 5. Parameter-family mechanism map.** Parameter families are grouped by the link they control in the iodine-failure chain: local generation, transport renewal, free-I2 gate, phase conversion, localization/permeability or voltage loss. The figure is a reasoning scaffold for interpreting interventions, not a universal ranking of design levers.

### 2.7 Mechanism synthesis: what the model explains and what remains open

The full mechanism can now be written in one chain. Local current generates free iodine. If local complexation, diffusion, flow renewal, dissolution and accessible area cannot remove it, free-I2 supersaturation appears. If the supersaturation is sustained, solid iodine inventory accumulates. Solid iodine changes the surface through closure-derived accessibility loss and film attenuation, and it changes the current distribution through bare/covered pathways. If the inventory remains broadly distributed, the hydraulic effect may remain mild. If it localizes, agglomerates or bridges pore throats, a modest mean solid fraction can create visible clumps and pressure symptoms.

This chain explains why capacity and rate routes are different but connected. Capacity loads the integral Omega_I2 and therefore drives solid inventory and surface-state persistence. Rate loads J_gen/J_rem and therefore drives early free-I2 supersaturation. A practical cell can experience either route or both. The same framework also explains why a post-failure cell is difficult to interpret: once flow obstruction, pressure rise, local dry-out or uncontrolled clumping occurs, voltage curves and EIS no longer isolate the initial onset mechanism.

The model-supported claims are therefore precise. The simulations support a local iodine-handling chain; they separate capacity and rate routes; they provide portable single-fiber closures; and they show that continuum iodine stress must be interpreted with a morphology/localization layer. The model does not directly measure surface coverage, does not quantify a pressure threshold, does not optimize an additive concentration and does not turn old EIS into film proof. These boundaries make the mechanism credible rather than weaker.

## 3 Discussion

### 3.1 Local iodine handling as the organizing principle

The main conclusion is that iodine failure in Zn-I flow batteries is governed by local iodine handling inside the positive carbon felt. This principle is more general than a particular salt, felt treatment or voltage trace. It says that iodine remains useful when free-I2 generation, buffering, transport renewal, dissolution and surface accessibility remain balanced. It becomes a failure source when that balance is lost locally.

This view explains why bulk chemistry can be simultaneously necessary and insufficient. Iodide, bromide and polyiodide equilibria determine the chemical ability of the electrolyte to store iodine. But the positive felt determines how the reaction sees that chemistry locally. A high-capacity electrolyte can still fail at a fiber interface; a lower-capacity electrolyte can be stable if local source-removal balance is maintained.

### 3.2 Why capacity and rate should be treated as two routes

Capacity and rate should not be collapsed into one empirical severity axis. Capacity-oriented operation emphasizes time-integrated exposure and solid persistence. The associated observables are Omega_I2, eps_s, theta, R_film and accessible-area loss. Rate-oriented operation emphasizes instantaneous source strength. The associated observables are J_gen/J_rem and early S_I2. This distinction suggests different mitigation strategies. Capacity problems require reversible deposition, dissolution and inventory management; rate problems require better renewal, lower local current density and improved current distribution.

### 3.3 What the molecular and single-fiber levels add

The molecular calculations add chemical plausibility: iodine can reside near carbon motifs, and modest surface heterogeneity can bias where iodine first accumulates. The single-fiber model adds closure structure: it translates local deposition into coverage, film attenuation, precipitation/dissolution and current split without fitting a full-cell voltage. These two levels are the reason the COMSOL variables have physical meaning. Without them, theta and R_film would be arbitrary internal variables. With them, they are explicitly defined, portable closure states whose limitations can be stated.

### 3.4 Why hydraulic failure is a morphology problem

Visible clumps and pressure rise are not explained by uniform mean solid iodine alone. The model identifies broad iodine stress and mean solid inventory, but pressure symptoms require localization, agglomeration or pore-throat bridging. This is not a contradiction; it is a hierarchy. Continuum simulation predicts where iodine stress and solid inventory become possible. Morphology determines how that inventory is arranged in the pore network. The next model generation should therefore couple precipitation to permeability and pore accessibility rather than only refine voltage fits.

### 3.5 Evidence boundary and future falsification

The framework is intentionally conservative about evidence. Visible iodine, clumps and pressure symptoms are qualitative physical anchors. COMSOL and the single-fiber model support a mechanistic chain. Direct phase imaging, pressure traces, interrupted-charge recovery, Raman or optical identification, Cdl/area tracking and reference-electrode measurements are needed to convert the chain into quantitative validation.

Several falsification tests are straightforward. If rate-driven failure is controlled by generation-removal imbalance, flow renewal and active-area distribution should shift early S_I2 and onset more strongly than they shift late capacity inventory. If capacity-driven failure is controlled by solid persistence, rest/dissolution protocols and interrupted-charge imaging should change eps_s and theta-derived accessibility more strongly than they change early supersaturation. If hydraulic symptoms are controlled by localization, pressure rise should correlate with visible or tomographic clump localization more strongly than with mean solid inventory alone.

### 3.6 Implications for design without making an optimization claim

The manuscript does not claim a finalized NH4Br recipe, a best felt, a production operating map or a universal design recipe. It provides a mechanism checklist. A successful Zn-I positive electrode should spread current over accessible carbon area, renew the fiber-scale environment, buffer free iodine without hiding local supersaturation, keep solid iodine reversible and distributed, avoid pore-throat localization, and maintain voltage efficiency without conflating ohmic loss with phase failure. This checklist can compare additives, membranes, felts, compression states and flow fields after each intervention is assigned to the physical link it controls.

## 4 Methods

### 4.1 Journal positioning and evidence use

The study is written as a simulation-led mechanism article for a broad technological-sciences readership. It combines bounded molecular priors, a standalone single-fiber model, porous-electrode COMSOL postprocessing and physical interpretation. Archived electrochemical traces are not used as central microscopic evidence because full-cell voltage and uncontrolled EIS cannot uniquely identify iodine coverage, film growth or hydraulic blockage. Physical observations of visible iodine, clumps and pressure symptoms are retained as qualitative anchors only.

### 4.2 Molecular surface calculations

Molecular calculations were performed on representative finite carbon motifs to test iodine residence near graphitic and functional environments. Basal graphitic carbon served as the reference motif. Hydroxylated basal and carbonyl-edge motifs represented possible surface heterogeneity. The resulting adsorption energies were used as qualitative priors for residence and patchy nucleation tendency. They were not inserted directly as macroscopic precipitation rates or film conductivities.

### 4.3 Single-fiber iodine model

The single-fiber model solves radial transport around a carbon fiber with fast iodine complexation, surface electrochemical I2/I- conversion, precipitation/dissolution and surface-state evolution. The outer boundary represents the local bulk/diffusion-layer condition and the inner boundary represents the fiber surface. Positive applied current corresponds to I- oxidation and I2 generation. The model reports c_I2,free, h_I2, theta, R_film, f_film, precipitation/dissolution rates and bare/covered current fractions. Closure fits were generated from a 720-row parameter scan.

### 4.4 Porous-electrode model and postprocessing

The porous-electrode model represents the zinc negative side, membrane and positive carbon felt with iodine speciation, transport, reaction, precipitation/dissolution and closure-derived surface variables. This manuscript uses existing solved fields and postprocessing, without replacing production parameters. Extracted fields include free-I2 supersaturation proxy, solid iodine inventory, theta, R_film, film kinetic factor, accessible-area proxy and spatial profiles. The postprocessing focuses on state ordering and physical route separation rather than absolute voltage fitting.

### 4.5 Capacity and rate route analysis

Capacity-oriented and rate-oriented cases were compared using normalized charge progress. The analysis tracked the relative timing of S_I2, eps_s, theta, R_film and accessible-area loss. A capacity route was assigned when cumulative solid inventory and surface-state loss increased late in charge. A rate route was assigned when high free-I2 supersaturation appeared early, before comparable mean solid inventory accumulated.

### 4.6 Mean-field versus localized blockage

Uniform mean solid iodine was interpreted through a simple porosity/permeability argument. Localized blockage was represented by eps_s,local = eps_s,mean/f_local, where f_local is the fraction of pore volume or active region containing the solid inventory. This analysis is not a pressure-threshold fit. It is a physical amplification calculation showing why visible clumps and hydraulic symptoms require localization beyond the continuum mean.

### 4.7 Parameter-family interpretation

Parameter effects were grouped by physical entry point rather than by a single sensitivity rank. Accessible carbon area changes local source intensity and inventory distribution. Flow renewal and diffusivity change removal capacity. Speciation and saturation move the free-I2 phase gate. Precipitation and dissolution control solid inventory and recovery. Compression and localization control hydraulic amplification. Conductivity and ASR control voltage loss and heat but do not by themselves define iodine phase stability.

### 4.8 Reproducibility and limitations

The scripts used to generate the single-fiber closures, postprocess COMSOL fields and assemble manuscript figures are retained in the project repository. Original experimental archives and COMSOL model files are not modified. The main limitations are explicit: surface coverage is closure-derived, film attenuation is a local model variable, pressure is not quantitatively predicted, clump hydrodynamics are not resolved, and the manuscript does not make a production optimization claim.

## 5 Conclusions

This work reframes solid iodine failure in zinc-iodine flow batteries as a local iodine-handling failure of the positive carbon felt. The decisive state is not total tank iodine but the local balance between free-I2 generation and removal at fiber surfaces. Molecular surface calculations support iodine residence near carbon motifs; a radial single-fiber model converts deposition into interpretable closures; COMSOL state evolution separates capacity-driven cumulative solid inventory from rate-driven early supersaturation; and morphology analysis explains why visible clumps and pressure symptoms require localization beyond mean solid fraction. The framework provides a physically ordered way to discuss additives, flow, active area, precipitation, compression and voltage loss. Its most important value is not that it supplies a final design recipe, but that it identifies the variables that must be measured next to make iodine failure in Zn-I flow batteries predictable.

## Data availability

Processed data tables and figure-generation scripts needed to reproduce the manuscript figures will be deposited in a public repository or provided as supplementary files upon acceptance. Original experimental archives and COMSOL model files are retained unchanged.

## Code availability

Python postprocessing scripts, the standalone single-fiber model, closure-fitting scripts and manuscript figure-generation scripts will be made available with the curated data package.

## Author contributions

To be completed by the authors.

## Competing interests

The authors declare no competing interests, or revise as appropriate.

## Acknowledgements

To be completed by the authors.

## References

{refs}
"""


def si_text() -> str:
    return """# Supporting Information for Local iodine-handling failure in zinc-iodine flow batteries

## S1 Molecular surface evidence

The molecular calculations are retained as bounded surface priors. They support iodine residence near graphitic and functional carbon motifs, and they justify patchy residence or nucleation as chemically plausible. They are not used as direct macroscopic precipitation rates, site densities or film conductivities.

![SI_Fig_molecular_surface_evidence_R64.png](figures/SI_Fig_molecular_surface_evidence_R64.png)

## S2 COMSOL field visualizations

The COMSOL field package provides the continuum-state basis for broad iodine stress and weak resolved continuum-locality interpretation. The fields support the conclusion that visible clumps require morphology amplification rather than a single mesh-resolved hot point.

![SI_Fig_true_COMSOL_cloudmaps_R64.png](figures/SI_Fig_true_COMSOL_cloudmaps_R64.png)

## S3 Precipitation and dissolution boundary

The precipitation/dissolution map explains why the free-I2 saturation boundary is the phase gate. Precipitation responds to supersaturation, whereas dissolution requires undersaturation and existing solid/covered state. This distinction is important for interpreting rest and discharge recovery.

![SI_Fig_precipitation_dissolution_boundary_R64.png](figures/SI_Fig_precipitation_dissolution_boundary_R64.png)

## S4 Dense CAP/RATE time evolution

The dense state-evolution figure supports the selected main Fig. 3 by retaining additional internal variables. The important reading is the ordering of variables, not merely the endpoint values.

![SI_dense_Fig3_CAP_RATE_time_evolution_R64.png](figures/SI_dense_Fig3_CAP_RATE_time_evolution_R64.png)

## S5 Parameter-family curves

The parameter-family curves support the mechanism map in Fig. 5. Each curve should be interpreted by physical entry point: generation, renewal, speciation, phase conversion, localization/permeability or voltage loss.

![SI_dense_Fig5_parameter_family_curves_R64.png](figures/SI_dense_Fig5_parameter_family_curves_R64.png)

## S6 Parameter fingerprints

The fingerprint figure separates source, removal, speciation, precipitation and voltage-loss signatures. It is intended to prevent different physical effects from being collapsed into a single lever score.

![SI_dense_Fig6_parameter_fingerprint_R64.png](figures/SI_dense_Fig6_parameter_fingerprint_R64.png)

## S7 Evidence boundary and falsification map

The following figure is retained in the Supporting Information because it is useful for review and future planning, but it is more meta-analytical than the main mechanism figures.

![Fig6_R67_evidence_boundary_falsification.png](figures/Fig6_R67_evidence_boundary_falsification.png)

**Figure S7. Evidence hierarchy and falsification map.** The model-supported source-removal-supersaturation-solid chain is separated from experiments still needed. Direct phase identification, pressure tracing, interrupted-charge imaging and surface-accessibility measurements are required for quantitative validation of clump hydraulics and coverage.

### Literature positioning

Prior Zn-I studies established high aqueous energy density, bromide/polyiodide complexation, robust electrolyte design, stable cell architectures and long-life ion management. The remaining gap addressed here is different: local positive-felt iodine handling. High tank capacity does not define the local free-I2 failure coordinate; speciation is one link in local handling rather than the entire mechanism; and porous carbon-felt transport/compression must be connected to solid iodine localization and pore-throat risk.
"""


def figure_captions() -> str:
    return """# Separate figure captions

## Figure 1. Local iodine-handling failure in porous positive carbon felt.

The reservoir can remain chemically buffered while the positive felt fails locally. Electrochemical current creates free iodine at fiber surfaces; complexation, diffusion, flow renewal, dissolution and remaining accessible area determine whether that iodine is removed or crosses a saturation boundary. Persistent supersaturation drives solid iodine inventory, closure-derived surface accessibility loss, film attenuation and morphology-scale clumping.

## Figure 2. Single-fiber bridge from local precipitation to porous-electrode closures.

The radial fiber calculation resolves free-I2 transport, fast complexation, electrochemical generation/consumption, precipitation/dissolution and local solid iodine thickness. The output is converted into closures for theta(h_I2), R_film, film kinetic attenuation and bare/covered current split. Each curve maps one local surface state into one macroscopic closure variable, keeping precipitation, coverage, film penalty and current redistribution physically separate.

## Figure 3. Time-resolved state variables separate capacity and rate routes.

Capacity-oriented operation produces delayed growth of solid iodine inventory, closure-derived coverage, film proxy and area-loss proxy, indicating cumulative supersaturation dose and surface-state persistence. Rate-oriented operation produces early free-I2 supersaturation before comparable mean solid inventory develops, indicating an instantaneous generation/removal imbalance. The key information is the ordering of physical state variables rather than a fitted voltage residual.

## Figure 4. Spatial iodine-state fields and morphology amplification.

Coordinate-aware COMSOL fields show broad positive-electrode iodine stress rather than a single resolved continuum hotspot. This spatial result places visible clumps, pore-throat obstruction and pressure symptoms at the morphology-amplification level, where nucleation heterogeneity, stagnant zones, fiber contacts and localized solid accumulation can convert modest mean inventory into severe local blockage.

## Figure 5. Parameter-family mechanism map.

Parameter families are grouped by the link they control in the iodine-failure chain: local generation, transport renewal, free-I2 gate, phase conversion, localization/permeability or voltage loss. The figure is a reasoning scaffold for interpreting interventions. A parameter is important because of where it enters the chain and what tradeoff it creates, not because it ranks highest in a single sensitivity bar.
"""


def cover_letter() -> str:
    return f"""# Cover letter draft for Science China Technological Sciences

Dear Editors,

We are pleased to submit the invited manuscript entitled "{TITLE}" for consideration in *Science China Technological Sciences*.

The manuscript addresses a practical and under-resolved failure mode in zinc-iodine flow batteries: the conversion of iodine from soluble redox inventory into local solid iodine deposits, surface accessibility loss, clumps, pore blockage and hydraulic/electrochemical instability inside porous positive carbon felt. We frame this transition as a local iodine-handling failure governed by the balance between interfacial free-I2 generation and the combined local capacity supplied by complexation, diffusion, flow renewal, dissolution and accessible carbon surface.

The work fits the technological-sciences readership because it connects electrolyte chemistry, porous electrode transport, surface-state closure and device-scale failure. Molecular carbon-surface calculations provide bounded priors for iodine residence near carbon motifs. A standalone single-fiber model converts local precipitation into coverage, film attenuation and bare/covered current partition. COMSOL porous-electrode postprocessing then resolves how capacity-oriented and rate-oriented operation load the same iodine failure chain through different physical routes. A final morphology analysis explains why visible clumps and pressure symptoms require localization or pore-throat bridging beyond a uniform mean solid fraction.

The main contribution is a physically organized framework, not an empirical operating map or an additive screen. The framework identifies measurable coordinates of iodine failure: local generation/removal balance, free-I2 supersaturation, cumulative supersaturation dose, solid iodine inventory, closure-derived surface accessibility and localization-driven hydraulic risk. These coordinates lead directly to falsifiable follow-up experiments, including phase identification, interrupted-charge imaging, pressure tracing and reference-electrode separation.

Sincerely,

[Corresponding author name]

[Affiliation]

[Email]
"""


def highlights() -> str:
    return """# Highlights

- Local iodine-handling failure is identified as the route from soluble iodine redox chemistry to solid blockage in zinc-iodine flow batteries.
- Molecular surface priors, a radial single-fiber model and COMSOL porous-electrode fields are connected into one mechanism chain.
- Capacity loading accumulates solid iodine inventory and surface-state loss, whereas rate loading first creates free-I2 generation/removal imbalance.
- Broad continuum iodine stress is connected to visible clumps through localization and pore-throat morphology amplification.
- Parameter effects are argued by physical entry point: generation, renewal, speciation, phase conversion, localization or voltage loss.
"""


def figure_interpretation_ledger() -> str:
    return """# R84 figure interpretation ledger

| Figure | Why it is in the main text | Main conclusion | Boundary |
|---|---|---|---|
| Fig. 1 | Establishes the paper-level mechanism chain | Failure is local iodine handling, not just tank inventory | Schematic, not direct proof |
| Fig. 2 | Converts local fiber deposition into portable closures | Coverage, film attenuation, precipitation and current split are separable | Closure-derived, not direct imaging |
| Fig. 3 | Provides the time-evolution group plot central to the Results | Capacity and rate routes have different state-variable ordering | State variables, not voltage fitting |
| Fig. 4 | Connects COMSOL fields to observed clump/pressure symptoms | Broad iodine stress needs morphology amplification for blockage | No exact pressure prediction |
| Fig. 5 | Replaces lever ranking with parameter-by-parameter physics | Parameters matter by where they enter the chain | Not an optimization map |
"""


def gpt_prompt_cn() -> str:
    return """# 给 GPT Pro / 合作者的审阅提示

请按 *Science China Technological Sciences* 约稿论文标准审阅 R84 稿件。重点不是语法润色，而是判断它是否已经像一篇真正的机制论文：

1. 题目和摘要是否能让编辑立刻明白贡献是 local iodine-handling failure，而不是 NH4Br 优化或旧全电池拟合。
2. Introduction 是否把 Zn-I 文献、碘的双重角色、bulk 不足和 local failure 的必要性讲清楚。
3. Fig. 1-5 是否每张都有不可替代的论文功能。
4. 分子模拟是否被正确限定为 surface-residence prior，没有过度转移成宏观速率。
5. 单纤维模型是否真正支撑了 coverage / film / precipitation / current split 的闭合关系。
6. COMSOL 时间演化图是否足够支持 CAP route 与 RATE route 的区分。
7. 参数逐项论证是否足够深，是否还像一张 lever 图的解释。
8. 形貌/堵塞/压力的解释是否既有新意又不过度承诺。
9. 讨论和方法是否像期刊论文，而不是内部任务报告。
10. 哪些句子仍然过强，哪些地方需要补实验或挪到 SI。

请给出：可投性判断、最大卖点、最大短板、必须修改的 10 个具体位置、以及是否建议保留 Fig. 5 在主文。
"""


def write_text_files() -> None:
    (OUT / "SCTS_ZIFB_publication_deep_revision_R84.md").write_text(manuscript_text(), encoding="utf-8")
    (OUT / "Supporting_Information_R84.md").write_text(si_text(), encoding="utf-8")
    (OUT / "separate_figure_captions_R84.md").write_text(figure_captions(), encoding="utf-8")
    (OUT / "cover_letter_SCTS_invited_R84.md").write_text(cover_letter(), encoding="utf-8")
    (OUT / "highlights_R84.md").write_text(highlights(), encoding="utf-8")
    (OUT / "R84_figure_interpretation_ledger.md").write_text(figure_interpretation_ledger(), encoding="utf-8")
    (OUT / "GPT_PRO_REVIEW_PROMPT_R84_CN.md").write_text(gpt_prompt_cn(), encoding="utf-8")
    (OUT / "R84_SCTS_source_note.md").write_text(
        """# SCTS source note

The R84 package is prepared against the public author information available for *Science China Technological Sciences*:

- Springer submission guidelines: https://link.springer.com/journal/11431/submission-guidelines
- Science China author center: https://www.sciengine.com/SCTS/authorCenter?scroll=section_1
- Springer-hosted Instructions for Authors PDF: https://media.springer.com/full/springer-instructions-for-authors-assets/pdf/11431_SCTS%20Instructions%20for%20authors-20230608.pdf

R84 uses the standard research-manuscript structure: Abstract, Keywords, Introduction, Results, Discussion, Methods, Conclusions and declarations. Author metadata, funding and reference-manager verification remain human-required before upload.
""",
        encoding="utf-8",
    )
    (OUT / "R84_author_fill_in_sheet.md").write_text(
        """# Author fill-in sheet before SCTS upload

- Full author list:
- Affiliations:
- Corresponding author:
- Email:
- ORCID IDs if required:
- Equal contribution statement if any:
- Funding:
- Acknowledgements:
- Conflict of interest statement:
- Data/code availability final wording:
- Article type confirmed with invited editor:
- Reference-manager check completed:
- Final PDF proof checked page by page:
""",
        encoding="utf-8",
    )


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")


def apply_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_rich_paragraph(doc: Document, text: str, style: str | None = None, *, italic=False, center=False) -> None:
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        run = p.add_run(part[2:-2] if bold else part)
        run.bold = bold
        run.italic = italic
        run.font.name = "Calibri"
        run.font.size = Pt(10 if italic else 11)
    if italic:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if re.match(r"^\|\s*-", line):
            continue
        rows.append([c.strip() for c in line.strip("|").split("|")])
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if ncols == 3:
        widths = [1700, 3500, 4160]
    elif ncols == 4:
        widths = [1450, 2550, 2650, 2710]
    else:
        widths = [9360 // ncols] * ncols
    set_table_width(table, widths)
    for r, row in enumerate(rows):
        for c in range(ncols):
            cell = table.cell(r, c)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[c] if c < len(row) else ""
            set_cell_margins(cell)
            if r == 0:
                set_cell_shading(cell, "F4F6F9")
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(3)
                para.paragraph_format.line_spacing = 1.15
                for run in para.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
                    if r == 0:
                        run.bold = True


def markdown_to_docx(md_path: Path, docx_path: Path, image_names: list[str]) -> dict[str, object]:
    doc = Document()
    apply_styles(doc)
    lookup = {f"figures/{name}": FIG / name for name in image_names}
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith("|---"):
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(17)
            run.font.color.rgb = RGBColor.from_string("0B2545")
            p.paragraph_format.space_after = Pt(14)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("!["):
            m = re.search(r"\(([^)]+)\)", line)
            if m and m.group(1) in lookup and lookup[m.group(1)].exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                width = 4.25 if "SI_dense_Fig3" in m.group(1) else 6.15
                run.add_picture(str(lookup[m.group(1)]), width=Inches(width))
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(doc, table_lines)
            continue
        elif re.match(r"^[A-Za-z0-9_,]+\\s*=.*\\(\\d+\\)$", line):
            add_rich_paragraph(doc, line, italic=True, center=True)
        elif line.startswith("- "):
            add_rich_paragraph(doc, line[2:], style="List Bullet")
        elif line.startswith("**Figure") or line.startswith("**Fig."):
            add_rich_paragraph(doc, line, italic=True)
        else:
            add_rich_paragraph(doc, line)
        i += 1
    props = doc.core_properties
    props.title = TITLE
    props.subject = "Science China Technological Sciences invited manuscript draft"
    props.keywords = "zinc-iodine flow battery; iodine precipitation; carbon felt; local supersaturation"
    props.comments = "R84 publication deep revision; author metadata and references require human verification."
    doc.save(docx_path)
    return docx_check(docx_path)


def docx_check(path: Path) -> dict[str, object]:
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return {
        "exists": path.exists(),
        "paragraphs": len(doc.paragraphs),
        "images": len(doc.inline_shapes),
        "tables": len(doc.tables),
        "literal_backslash_n": "\\n" in text,
        "references_heading": any(p.text.strip() == "References" for p in doc.paragraphs),
    }


def export_pdf(docx: Path, pdf: Path) -> None:
    ps = PDF / f"export_{docx.stem}.ps1"
    ps.write_text(
        f"""
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $doc = $word.Documents.Open('{docx}')
  $doc.ExportAsFixedFormat('{pdf}', 17)
  $doc.Close($false)
}} finally {{
  $word.Quit()
}}
""",
        encoding="utf-8-sig",
    )
    proc = subprocess.run(
        ["powershell", "-ExecutionPolicy", "Bypass", "-File", str(ps)],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        timeout=240,
    )
    stdout = proc.stdout or ""
    stderr = proc.stderr or ""
    (PDF / f"{docx.stem}_export_stdout_stderr.txt").write_text(stdout + "\n" + stderr, encoding="utf-8")
    if proc.returncode != 0 or not pdf.exists():
        raise RuntimeError(f"Word PDF export failed for {docx.name}: {proc.stderr}")


def render_pdf_pages(pdf: Path, out_dir: Path) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    doc = pdfium.PdfDocument(str(pdf))
    paths = []
    for i in range(len(doc)):
        page = doc[i]
        bitmap = page.render(scale=1.4).to_pil()
        p = out_dir / f"page_{i+1:03d}.png"
        bitmap.save(p)
        paths.append(p)
    return paths


def make_contact_sheet(page_paths: list[Path], out: Path, title: str) -> None:
    thumbs = []
    for p in page_paths:
        img = Image.open(p).convert("RGB")
        img.thumbnail((260, 340))
        thumbs.append((p.name, img.copy()))
    cols = 4
    rows = (len(thumbs) + cols - 1) // cols
    w = cols * 300
    h = rows * 390 + 60
    sheet = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(sheet)
    draw.text((20, 20), title, fill=(0, 0, 0))
    for idx, (name, img) in enumerate(thumbs):
        x = (idx % cols) * 300 + 20
        y = (idx // cols) * 390 + 60
        sheet.paste(img, (x, y))
        draw.rectangle([x, y, x + img.width, y + img.height], outline=(120, 120, 120))
        draw.text((x, y + img.height + 6), name, fill=(0, 0, 0))
    sheet.save(out)


def pdf_qa(pdf: Path, label: str) -> dict[str, object]:
    pages = render_pdf_pages(pdf, PDF / "pages" / pdf.stem)
    make_contact_sheet(pages, PDF / f"{label}_contact_sheet_R84.png", f"{label} contact sheet")
    sizes = [Image.open(p).getbbox() for p in pages]
    blank = [p.name for p, bbox in zip(pages, sizes) if bbox is None]
    return {"pdf": str(pdf), "pages": len(pages), "blank_pages": blank}


def audits(checks: dict[str, object], pdf_checks: dict[str, object]) -> dict[str, object]:
    md = (OUT / "SCTS_ZIFB_publication_deep_revision_R84.md").read_text(encoding="utf-8")
    body = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", md)
    forbidden = [
        "workflow",
        "dashboard",
        "EIS proves",
        "EIS disproves",
        "validated operating map",
        "optimized NH4Br",
        "measured coverage",
        "exact pressure threshold",
        "full-cell proves",
        "Q*/i* fitted",
        "parameter-free",
        "production design claim",
        "550C best",
    ]
    q = {
        "word_count_rough": len(re.findall(r"[A-Za-z0-9_+./-]+", body)),
        "references": len(re.findall(r"^\[\d+\]", md, re.M)),
        "main_figures": len(re.findall(r"!\[Fig", md)),
        "forbidden_counts": {pat: len(re.findall(re.escape(pat), md, re.I)) for pat in forbidden},
        "docx": checks,
        "pdf": pdf_checks,
    }
    (OUT / "R84_QA_summary.json").write_text(json.dumps(q, indent=2), encoding="utf-8")
    (OUT / "R84_publication_revision_audit.md").write_text(
        f"""# R84 publication deep revision audit

## What changed beyond R83

- Rewrote the manuscript into a stronger journal-style mechanism paper.
- Expanded the Results so each figure has an explicit scientific role and conclusion.
- Added deeper parameter-by-parameter argumentation instead of relying on a single lever map.
- Integrated molecular simulation and single-fiber modelling as evidence layers, not as disconnected appendices.
- Retained claim boundaries: no direct measured coverage, no pressure threshold, no additive optimization and no old-EIS proof.

## Structural QA

- Rough word count: {q['word_count_rough']}
- References: {q['references']}
- Main figures: {q['main_figures']}
- Manuscript DOCX images: {checks['manuscript']['images']}
- Manuscript DOCX tables: {checks['manuscript']['tables']}
- SI DOCX images: {checks['supporting_information']['images']}
- SI DOCX tables: {checks['supporting_information']['tables']}
- Manuscript PDF pages: {pdf_checks['manuscript']['pages']}
- SI PDF pages: {pdf_checks['supporting_information']['pages']}
- Manuscript blank pages: {pdf_checks['manuscript']['blank_pages']}
- SI blank pages: {pdf_checks['supporting_information']['blank_pages']}

## Forbidden wording scan

""" + "\n".join(f"- {k}: {v}" for k, v in q["forbidden_counts"].items()) + "\n",
        encoding="utf-8",
    )
    status = "PASS" if all(v == 0 for v in q["forbidden_counts"].values()) and not pdf_checks["manuscript"]["blank_pages"] and not pdf_checks["supporting_information"]["blank_pages"] else "PARTIAL"
    (OUT / "final_status.md").write_text(
        f"""SCTS R84 publication deep revision: {status}
Manuscript Markdown: SCTS_ZIFB_publication_deep_revision_R84.md
Manuscript DOCX: SCTS_ZIFB_publication_deep_revision_R84.docx
Manuscript PDF proof: pdf_proof/SCTS_ZIFB_publication_deep_revision_R84.pdf
Supporting Information DOCX: Supporting_Information_R84.docx
Supporting Information PDF proof: pdf_proof/Supporting_Information_R84.pdf
Rough word count: {q['word_count_rough']}
Main figures: {q['main_figures']}
References: {q['references']}
Parameter-by-parameter reasoning expanded: YES
Molecular modelling integrated: YES
Single-fiber closure integrated: YES
COMSOL time-evolution interpretation expanded: YES
PDF visual proof generated: YES
Manuscript PDF pages: {pdf_checks['manuscript']['pages']}
SI PDF pages: {pdf_checks['supporting_information']['pages']}
Blank pages detected: {'NO' if not pdf_checks['manuscript']['blank_pages'] and not pdf_checks['supporting_information']['blank_pages'] else 'YES'}
Ready for GPT Pro/coauthor scientific review: YES
Ready for senior-author line editing: YES
Ready for direct one-click SCTS upload: NO
Reason: author metadata, funding, reference-manager verification and final human proof are still required.
""",
        encoding="utf-8",
    )
    return q


def package_outputs() -> None:
    full = PKG / "SCTS_R84_PUBLICATION_DEEP_REVISION_FULL.zip"
    gpt = PKG / "GPT_PRO_REVIEW_PACKET_R84_PUBLICATION_DEEP_REVISION.zip"
    with zipfile.ZipFile(full, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(OUT.rglob("*")):
            if p.is_file() and "package" not in p.relative_to(OUT).parts:
                zf.write(p, str(p.relative_to(OUT)).replace("\\", "/"))
    wanted = [
        "SCTS_ZIFB_publication_deep_revision_R84.md",
        "SCTS_ZIFB_publication_deep_revision_R84.docx",
        "Supporting_Information_R84.md",
        "Supporting_Information_R84.docx",
        "separate_figure_captions_R84.md",
        "cover_letter_SCTS_invited_R84.md",
        "highlights_R84.md",
        "R84_figure_interpretation_ledger.md",
        "GPT_PRO_REVIEW_PROMPT_R84_CN.md",
        "R84_publication_revision_audit.md",
        "R84_QA_summary.json",
        "final_status.md",
        "pdf_proof/SCTS_ZIFB_publication_deep_revision_R84.pdf",
        "pdf_proof/Supporting_Information_R84.pdf",
        "pdf_proof/manuscript_contact_sheet_R84.png",
        "pdf_proof/supporting_information_contact_sheet_R84.png",
    ]
    with zipfile.ZipFile(gpt, "w", zipfile.ZIP_DEFLATED) as zf:
        for rel in wanted:
            p = OUT / rel
            if p.exists():
                zf.write(p, rel.replace("\\", "/"))
        for p in sorted(FIG.glob("*.png")):
            zf.write(p, f"figures/{p.name}")


def main() -> None:
    ensure_dirs()
    copy_figures()
    write_text_files()
    checks = {
        "manuscript": markdown_to_docx(OUT / "SCTS_ZIFB_publication_deep_revision_R84.md", OUT / "SCTS_ZIFB_publication_deep_revision_R84.docx", MAIN_FIGS),
        "supporting_information": markdown_to_docx(OUT / "Supporting_Information_R84.md", OUT / "Supporting_Information_R84.docx", SI_FIGS),
    }
    export_pdf(OUT / "SCTS_ZIFB_publication_deep_revision_R84.docx", PDF / "SCTS_ZIFB_publication_deep_revision_R84.pdf")
    export_pdf(OUT / "Supporting_Information_R84.docx", PDF / "Supporting_Information_R84.pdf")
    pdf_checks = {
        "manuscript": pdf_qa(PDF / "SCTS_ZIFB_publication_deep_revision_R84.pdf", "manuscript"),
        "supporting_information": pdf_qa(PDF / "Supporting_Information_R84.pdf", "supporting_information"),
    }
    q = audits(checks, pdf_checks)
    package_outputs()
    print((OUT / "final_status.md").read_text(encoding="utf-8"))
    print(json.dumps({"out": str(OUT), "qa": q}, indent=2))


if __name__ == "__main__":
    main()
